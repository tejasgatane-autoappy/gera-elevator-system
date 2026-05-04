from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUTPUT = ROOT / "docs" / "Gera_Elevator_Frontend_Handoff.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text, bold=False):
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold
    run.font.name = "Arial"
    run.font.size = Pt(10)
    paragraph.paragraph_format.space_after = Pt(0)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def set_table_borders(table, color="D9DEE8"):
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:{}".format(edge)
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def add_heading(document, text, level=1):
    paragraph = document.add_heading(text, level=level)
    for run in paragraph.runs:
        run.font.name = "Arial"
        run.font.color.rgb = RGBColor(31, 52, 92)
    return paragraph


def add_bullet(document, text):
    paragraph = document.add_paragraph(style="List Bullet")
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_number(document, text):
    paragraph = document.add_paragraph(style="List Number")
    paragraph.paragraph_format.space_after = Pt(6)
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(11)


def add_code(document, text):
    paragraph = document.add_paragraph()
    paragraph.paragraph_format.left_indent = Inches(0.25)
    paragraph.paragraph_format.space_before = Pt(4)
    paragraph.paragraph_format.space_after = Pt(8)
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9.5)
    run.font.color.rgb = RGBColor(45, 55, 72)


def set_document_styles(document):
    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.08

    for style_name, size in (("Title", 22), ("Subtitle", 12), ("Heading 1", 16), ("Heading 2", 13)):
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        if style_name.startswith("Heading"):
            style.font.bold = True
            style.font.color.rgb = RGBColor(31, 52, 92)


def add_header_footer(document):
    section = document.sections[0]
    header = section.header
    paragraph = header.paragraphs[0]
    paragraph.text = "Gera Elevator System - Frontend Handoff"
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.runs[0].font.name = "Arial"
    paragraph.runs[0].font.size = Pt(9)
    paragraph.runs[0].font.color.rgb = RGBColor(100, 116, 139)

    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("Frontend Integration Brief")
    run.font.name = "Arial"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)


def add_table(document, headers, rows, widths):
    table = document.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for idx, header in enumerate(headers):
        set_cell_text(hdr_cells[idx], header, bold=True)
        set_cell_shading(hdr_cells[idx], "EEF3FB")
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            set_cell_text(cells[idx], value)
    for row in table.rows:
        for idx, width in enumerate(widths):
            row.cells[idx].width = Inches(width)
    document.add_paragraph()
    return table


def build_doc():
    document = Document()
    section = document.sections[0]
    section.top_margin = Inches(0.85)
    section.bottom_margin = Inches(0.85)
    section.left_margin = Inches(0.8)
    section.right_margin = Inches(0.8)

    set_document_styles(document)
    add_header_footer(document)

    title = document.add_paragraph(style="Title")
    title_run = title.add_run("Frontend Handoff: Gera Elevator System")
    title_run.font.name = "Arial"
    title_run.font.color.rgb = RGBColor(24, 43, 77)

    subtitle = document.add_paragraph(style="Subtitle")
    subtitle.add_run("From: Team Lead | To: Frontend Team Member | Purpose: Connect React UI with Spring Boot backend")

    add_heading(document, "Goal", level=1)
    document.add_paragraph(
        "The frontend must show the building floors, four elevators, current lift positions, moving direction, "
        "external Up/Down buttons, and inside-lift floor buttons. The backend is responsible for choosing the best "
        "elevator and returning the updated stop list."
    )

    add_heading(document, "How The Flow Works", level=1)
    for item in [
        "User clicks an external Up/Down button or an inside-lift floor button.",
        "React sends the request to the Spring Boot backend.",
        "Backend validates the request and reads current elevator state from Redis.",
        "Scheduler assigns the best elevator using same-direction, nearest-idle, then minimum-ETA rules.",
        "Backend saves the updated state and returns the assigned elevator, ETA, and updated stops.",
        "Frontend refreshes the grid and updates the selected lift's visual state.",
    ]:
        add_number(document, item)

    add_heading(document, "Base URL", level=1)
    document.add_paragraph("Use this URL while running locally:")
    add_code(document, "http://localhost:8080")

    add_heading(document, "API Endpoints", level=1)
    add_table(
        document,
        ["Purpose", "Method", "Endpoint", "Frontend Usage"],
        [
            ["Get all lift states", "GET", "/api/v1/elevators", "Call on page load and during refresh/polling."],
            ["External floor request", "POST", "/api/v1/requests", "Use for floor Up/Down buttons."],
            ["Internal lift request", "POST", "/api/v1/elevators/{id}/requests", "Use for inside buttons of lift A/B/C/D."],
            ["Recent backend events", "GET", "/api/v1/events?limit=20", "Optional debug/activity panel."],
            ["Reset system", "POST", "/api/v1/admin/reset", "Optional dev/admin button only."],
        ],
        [1.65, 0.85, 2.25, 2.55],
    )

    add_heading(document, "Request Examples", level=1)
    document.add_paragraph("External request when user presses Up on floor 4:")
    add_code(document, '{\n  "type": "EXTERNAL",\n  "floor": 4,\n  "direction": "UP"\n}')
    document.add_paragraph("Internal request when user inside lift A selects floor 9:")
    add_code(document, '{\n  "destinationFloor": 9\n}')

    add_heading(document, "Response Example", level=1)
    add_code(
        document,
        '{\n'
        '  "assignedElevator": "A",\n'
        '  "estimatedArrivalTime": 60,\n'
        '  "stopsUpdated": [4, 6],\n'
        '  "reason": "SAME_DIRECTION_PASSING",\n'
        '  "stateVersion": 2\n'
        '}',
    )

    add_heading(document, "State Response Shape", level=1)
    add_code(
        document,
        '{\n'
        '  "totalFloors": 16,\n'
        '  "version": 2,\n'
        '  "elevators": [\n'
        '    {\n'
        '      "id": "A",\n'
        '      "currentFloor": 4,\n'
        '      "direction": "UP",\n'
        '      "doorStatus": "CLOSED",\n'
        '      "stops": [6]\n'
        '    }\n'
        '  ]\n'
        '}',
    )

    document.add_section(WD_SECTION.NEW_PAGE)
    add_heading(document, "Frontend Mapping", level=1)
    add_table(
        document,
        ["Backend Field", "Meaning", "Frontend Display"],
        [
            ["id", "Elevator name A/B/C/D", "Use as lift column label."],
            ["currentFloor", "Current floor number", "Place elevator icon/block in that floor row."],
            ["direction", "UP, DOWN, or IDLE", "Show arrow up/down or neutral state."],
            ["doorStatus", "OPEN, CLOSED, OPENING, CLOSING", "Optional door animation/state color."],
            ["stops", "Pending floors", "Highlight upcoming stops or show inside panel active buttons."],
            ["estimatedArrivalTime", "Seconds until lift reaches request floor", "Optional toast/status text."],
        ],
        [1.55, 2.35, 3.4],
    )

    add_heading(document, "Suggested React Integration Steps", level=1)
    for item in [
        "Create an API helper file, for example src/api/elevatorApi.js.",
        "On app load, call GET /api/v1/elevators and store the response in React state.",
        "Render floors from totalFloors down to 1 so the top floor appears at the top.",
        "For each elevator, compare elevator.currentFloor with the floor row and render the lift block there.",
        "On external Up/Down click, POST to /api/v1/requests with type, floor, and direction.",
        "On inside button click, POST to /api/v1/elevators/{id}/requests with destinationFloor.",
        "After every successful request, refresh state using GET /api/v1/elevators.",
        "Add a simple polling interval, such as every 1000-2000 ms, until WebSocket/SSE is added later.",
    ]:
        add_number(document, item)

    add_heading(document, "Minimal API Helper", level=1)
    add_code(
        document,
        "const BASE_URL = 'http://localhost:8080/api/v1';\n\n"
        "export async function getElevators() {\n"
        "  const res = await fetch(`${BASE_URL}/elevators`);\n"
        "  if (!res.ok) throw new Error('Failed to load elevators');\n"
        "  return res.json();\n"
        "}\n\n"
        "export async function sendExternalRequest(floor, direction) {\n"
        "  const res = await fetch(`${BASE_URL}/requests`, {\n"
        "    method: 'POST',\n"
        "    headers: { 'Content-Type': 'application/json' },\n"
        "    body: JSON.stringify({ type: 'EXTERNAL', floor, direction })\n"
        "  });\n"
        "  if (!res.ok) throw new Error('External request failed');\n"
        "  return res.json();\n"
        "}\n\n"
        "export async function sendInternalRequest(elevatorId, destinationFloor) {\n"
        "  const res = await fetch(`${BASE_URL}/elevators/${elevatorId}/requests`, {\n"
        "    method: 'POST',\n"
        "    headers: { 'Content-Type': 'application/json' },\n"
        "    body: JSON.stringify({ destinationFloor })\n"
        "  });\n"
        "  if (!res.ok) throw new Error('Internal request failed');\n"
        "  return res.json();\n"
        "}",
    )

    add_heading(document, "Frontend Checklist", level=1)
    for item in [
        "Floor grid renders from 16 to 1.",
        "All four elevators A, B, C, D are visible.",
        "External Up/Down buttons call the backend with the correct floor and direction.",
        "Inside buttons for each lift call the correct elevator ID.",
        "UI refreshes after every successful backend response.",
        "Loading and error states are visible but simple.",
        "Invalid backend responses do not break the UI.",
        "Reset button is hidden from normal users or used only during development.",
    ]:
        add_bullet(document, item)

    add_heading(document, "Important Notes", level=1)
    add_bullet(document, "Do not calculate elevator assignment in the frontend. Backend is the source of truth.")
    add_bullet(document, "Use backend state to render elevator positions. Avoid maintaining separate fake lift state.")
    add_bullet(document, "For now, polling is acceptable. Later we can upgrade to WebSocket or Server-Sent Events.")
    add_bullet(document, "CORS is already enabled for localhost React dev ports 3000 and 5173.")

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    document.save(OUTPUT)


if __name__ == "__main__":
    build_doc()
    print(OUTPUT)

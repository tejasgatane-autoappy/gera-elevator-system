from pathlib import Path

from docx import Document


ROOT = Path(__file__).resolve().parents[1]
DOCX = ROOT / "docs" / "Gera_Elevator_Frontend_Handoff.docx"


def main():
    document = Document(DOCX)
    paragraphs = [p.text.strip() for p in document.paragraphs if p.text.strip()]
    required_phrases = [
        "Frontend Handoff: Gera Elevator System",
        "How The Flow Works",
        "API Endpoints",
        "Request Examples",
        "Frontend Mapping",
        "Suggested React Integration Steps",
        "Minimal API Helper",
        "Frontend Checklist",
    ]
    missing = [phrase for phrase in required_phrases if phrase not in paragraphs]
    if missing:
        raise SystemExit(f"Missing expected headings/content: {missing}")
    if len(document.tables) < 2:
        raise SystemExit("Expected at least two handoff tables")
    print(f"OK: {DOCX}")
    print(f"Paragraphs: {len(paragraphs)}")
    print(f"Tables: {len(document.tables)}")


if __name__ == "__main__":
    main()
